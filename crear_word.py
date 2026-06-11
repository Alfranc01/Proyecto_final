from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('Tablas de Frecuencia', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('Ejercicio 1: Empleados por sucursal (n=50)', level=1)

data1 = [
    ['Valor (xi)', 'Frecuencia absoluta (fi)', 'Frecuencia acumulada (Fi)', 'Frecuencia relativa (hi)', 'Frecuencia relativa acumulada (Hi)'],
    ['9', '2', '2', '4%', '4%'],
    ['10', '4', '6', '8%', '12%'],
    ['11', '10', '16', '20%', '32%'],
    ['12', '10', '26', '20%', '52%'],
    ['13', '5', '31', '10%', '62%'],
    ['14', '3', '34', '6%', '68%'],
    ['15', '6', '40', '12%', '80%'],
    ['16', '5', '45', '10%', '90%'],
    ['17', '3', '48', '6%', '96%'],
    ['18', '2', '50', '4%', '100%'],
    ['Total', '50', '', '100%', '']
]

table1 = doc.add_table(rows=11, cols=5)
table1.style = 'Table Grid'

for i in range(11):
    for j in range(5):
        table1.rows[i].cells[j].text = data1[i][j]
        if i == 0:
            table1.rows[i].cells[j].paragraphs[0].runs[0].bold = True

doc.add_heading('Ejercicio 2: Horas de estacionamiento (n=60)', level=1)

data2 = [
    ['Valor (xi)', 'Frecuencia absoluta (fi)', 'Frecuencia acumulada (Fi)', 'Frecuencia relativa (hi)', 'Frecuencia relativa acumulada (Hi)'],
    ['1', '5', '5', '8.3%', '8.3%'],
    ['2', '8', '13', '13.3%', '21.7%'],
    ['3', '12', '25', '20%', '41.7%'],
    ['4', '16', '41', '26.7%', '68.3%'],
    ['5', '9', '50', '15%', '83.3%'],
    ['6', '6', '56', '10%', '93.3%'],
    ['7', '4', '60', '6.7%', '100%'],
    ['Total', '60', '', '100%', '']
]

table2 = doc.add_table(rows=8, cols=5)
table2.style = 'Table Grid'

for i in range(8):
    for j in range(5):
        table2.rows[i].cells[j].text = data2[i][j]
        if i == 0:
            table2.rows[i].cells[j].paragraphs[0].runs[0].bold = True

doc.save('tabla_frecuencia.docx')
print("Word document created: tabla_frecuencia.docx")
